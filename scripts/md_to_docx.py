from pathlib import Path
import argparse
from copy import deepcopy
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from htmldocx import HtmlToDocx
from markdown import markdown


def _set_run_font(run, font_name: str, size_pt: int, bold: bool = False) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size_pt)
    run.bold = bold


def _add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Pagina ")
    _set_run_font(run, "Arial", 9)

    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def _style_document(document: Document, report_title: str) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.3)
    section.right_margin = Cm(2.3)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Arial"
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal_style.font.size = Pt(11)

    heading_1 = document.styles["Heading 1"]
    heading_1.font.name = "Arial"
    heading_1._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    heading_1.font.size = Pt(18)
    heading_1.font.bold = True
    heading_1.font.color.rgb = RGBColor(0x10, 0x34, 0x5C)

    heading_2 = document.styles["Heading 2"]
    heading_2.font.name = "Arial"
    heading_2._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    heading_2.font.size = Pt(14)
    heading_2.font.bold = True
    heading_2.font.color.rgb = RGBColor(0x1E, 0x5A, 0x8A)

    heading_3 = document.styles["Heading 3"]
    heading_3.font.name = "Arial"
    heading_3._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    heading_3.font.size = Pt(12)
    heading_3.font.bold = True
    heading_3.font.color.rgb = RGBColor(0x2E, 0x6C, 0x9D)

    header_paragraph = section.header.paragraphs[0]
    header_paragraph.text = report_title
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if header_paragraph.runs:
        _set_run_font(header_paragraph.runs[0], "Arial", 9, bold=True)
        header_paragraph.runs[0].font.color.rgb = RGBColor(0x57, 0x57, 0x57)

    footer_paragraph = section.footer.paragraphs[0]
    footer_paragraph.clear()
    _add_page_number(footer_paragraph)

    for paragraph in document.paragraphs:
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(8)
        paragraph.paragraph_format.line_spacing = 1.15

        if paragraph.style.name == "Heading 1":
            paragraph.paragraph_format.space_before = Pt(16)
            paragraph.paragraph_format.space_after = Pt(10)
        elif paragraph.style.name == "Heading 2":
            paragraph.paragraph_format.space_before = Pt(12)
            paragraph.paragraph_format.space_after = Pt(8)
        elif paragraph.style.name == "Heading 3":
            paragraph.paragraph_format.space_before = Pt(10)
            paragraph.paragraph_format.space_after = Pt(6)

    for table in document.tables:
        table.style = "Table Grid"
        if not table.rows:
            continue

        for cell in table.rows[0].cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "1E5A8A")
            tc_pr.append(shd)

            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    _set_run_font(run, "Arial", 10, bold=True)
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        for row in table.rows[1:]:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        _set_run_font(run, "Arial", 10)


def _add_toc(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'

    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")

    placeholder = OxmlElement("w:t")
    placeholder.text = "Haz clic derecho y selecciona Actualizar campo para generar el indice"

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_separate)
    run._r.append(placeholder)
    run._r.append(fld_char_end)


def _build_styled_document(content_document: Document, report_title: str) -> Document:
    document = Document()

    cover = document.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_run = cover.add_run("INFORME TECNICO")
    _set_run_font(cover_run, "Arial", 14, bold=True)
    cover_run.font.color.rgb = RGBColor(0x1E, 0x5A, 0x8A)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(report_title)
    _set_run_font(title_run, "Arial", 24, bold=True)
    title_run.font.color.rgb = RGBColor(0x10, 0x34, 0x5C)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("Documento generado automaticamente")
    _set_run_font(subtitle_run, "Arial", 11)
    subtitle_run.font.color.rgb = RGBColor(0x57, 0x57, 0x57)

    date_line = document.add_paragraph()
    date_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_line.add_run(datetime.now().strftime("%d/%m/%Y"))
    _set_run_font(date_run, "Arial", 11)
    date_run.font.color.rgb = RGBColor(0x57, 0x57, 0x57)

    document.add_page_break()

    toc_heading = document.add_paragraph("Tabla de contenido")
    toc_heading.style = "Heading 1"
    toc_paragraph = document.add_paragraph()
    _add_toc(toc_paragraph)

    document.add_page_break()

    body = document.element.body
    for element in content_document.element.body:
        if element.tag.endswith("sectPr"):
            continue
        body.append(deepcopy(element))

    return document


def _extract_report_title(markdown_text: str) -> str:
    report_title = "Informe"
    for line in markdown_text.splitlines():
        stripped = line.strip()
        if stripped.startswith("# "):
            report_title = stripped[2:].strip()
            break
    return report_title


def convert_markdown_to_docx(input_path: Path, output_path: Path | None = None) -> Path:
    if output_path is None:
        output_path = input_path.with_suffix(".docx")

    if not input_path.exists():
        raise FileNotFoundError(f"No existe el archivo de entrada: {input_path}")

    if input_path.suffix.lower() != ".md":
        raise ValueError(f"El archivo de entrada debe ser .md: {input_path}")

    if output_path.suffix.lower() != ".docx":
        raise ValueError(f"El archivo de salida debe ser .docx: {output_path}")

    markdown_text = input_path.read_text(encoding="utf-8")
    report_title = _extract_report_title(markdown_text)

    html = markdown(
        markdown_text,
        extensions=["extra", "tables", "fenced_code", "sane_lists"],
    )

    parser = HtmlToDocx()
    raw_document = parser.parse_html_string(html)
    document = _build_styled_document(raw_document, report_title)
    _style_document(document, report_title)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))

    return output_path


def _build_arg_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="Convierte Markdown a DOCX con estilo profesional por defecto."
    )
    parser.add_argument(
        "input",
        help="Ruta del archivo markdown de entrada (.md)",
    )
    parser.add_argument(
        "output",
        nargs="?",
        help="Ruta del archivo docx de salida (.docx). Si no se indica, se sobrescribe <input>.docx",
    )
    return parser


def main() -> int:
    parser = _build_arg_parser()
    args = parser.parse_args()

    input_path = Path(args.input)
    output_path = Path(args.output) if args.output else None

    try:
        final_output = convert_markdown_to_docx(input_path, output_path)
    except (FileNotFoundError, ValueError) as error:
        print(error)
        return 1

    print(f"DOCX generado: {final_output}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
